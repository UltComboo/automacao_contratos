# gerador.py - VERSÃO COM NEGRITO SELETIVO
from docx import Document
from docx.shared import Pt
import os
import re
from datetime import datetime
from config import SOCIEDADE, OUTORGADOS, MESES_PT, PALAVRAS_NEGRITO


class GeradorContratos:
    def __init__(self):
        pass

    def preparar_placeholders(self, dados_pessoa):
        """Prepara placeholders simples"""
        # CORREÇÃO DA DATA
        data_formatada = dados_pessoa.get('data_formatada', '')

        # Se a data já tem "Porto Alegre/RS", usar como está
        if "Porto Alegre" in data_formatada and "RS" in data_formatada:
            data_assinatura = data_formatada
        else:
            # Adicionar cidade/estado apenas se não estiver presente
            data_assinatura = f"Porto Alegre/RS, {data_formatada}"

        return {
            "{{NOME_COMPLETO}}": dados_pessoa.get('nome_completo', ''),
            "{{CPF}}": dados_pessoa.get('cpf', ''),
            "{{CPF_OUTORGANTE}}": dados_pessoa.get('cpf', ''),
            "{{ENDERECO_COMPLETO}}": dados_pessoa.get('endereco_completo', ''),
            "{{CIDADE_ESTADO}}": f"{dados_pessoa.get('cidade', '')}/{dados_pessoa.get('estado', '')}",
            "{{ENDERECO_CIDADE}}": dados_pessoa.get('cidade', ''),
            "{{ENDERECO_ESTADO}}": dados_pessoa.get('estado', ''),
            "{{OAB_NUMERO}}": dados_pessoa.get('oab_numero', ''),
            "{{OAB_UF}}": dados_pessoa.get('oab_uf', ''),
            "{{ESTADO_CIVIL}}": dados_pessoa.get('estado_civil', ''),
            "{{BRASILEIRO_A}}": "brasileiro",
            "{{RESIDENTE_A}}": "residente e domiciliado",
            "{{ADVOGADO_A}}": "advogado",
            "{{SOCIEDADE_NOME}}": SOCIEDADE['nome'],
            "{{SOCIEDADE_CNPJ}}": SOCIEDADE['cnpj'],
            "{{DATA_ASSINATURA}}": data_assinatura,
            "{{NOME_ASSINATURA}}": dados_pessoa.get('nome_completo', ''),
            "{{SEDE}}": SOCIEDADE['sede_cidade'],
        }

    def aplicar_negrito_seletivo(self, texto, nome_completo):
        """
        Aplica negrito APENAS nas palavras específicas, mantendo o resto normal.

        Args:
            texto: Texto completo do parágrafo (com placeholders já substituídos)
            nome_completo: Nome completo da pessoa

        Returns:
            Lista de tuplas [(texto1, negrito1), (texto2, negrito2), ...]
        """
        if not texto.strip():
            return [(texto, False)]

        # Lista de todas as palavras que devem ficar em negrito
        palavras_para_negrito = []

        # 1. Adicionar nome completo (se existir e estiver no texto)
        if nome_completo and nome_completo in texto:
            palavras_para_negrito.append(nome_completo)

        # 2. Adicionar palavras da lista PALAVRAS_NEGRITO que estão no texto
        for palavra in PALAVRAS_NEGRITO:
            if palavra in texto:
                palavras_para_negrito.append(palavra)

        # Se não há palavras para negrito, retorna tudo normal
        if not palavras_para_negrito:
            return [(texto, False)]

        # Criar padrão regex para encontrar todas as palavras
        # Ordenar por tamanho (maior primeiro) para evitar substituições parciais
        palavras_para_negrito.sort(key=len, reverse=True)
        padrao = '|'.join(re.escape(palavra) for palavra in palavras_para_negrito)

        # Dividir o texto pelas palavras que devem ser negritadas
        partes = re.split(f'({padrao})', texto)

        # Processar as partes
        runs = []
        for parte in partes:
            if parte:  # Ignorar strings vazias
                deve_negrito = parte in palavras_para_negrito
                runs.append((parte, deve_negrito))

        # Se nenhuma parte foi encontrada (pode acontecer com regex complexa)
        if not runs:
            runs = [(texto, False)]

        return runs

    def processar_paragrafo(self, paragraph, placeholders, nome_completo):
        """
        Processa um parágrafo: substitui placeholders e aplica negrito seletivo.
        """
        texto_original = paragraph.text

        # Pular parágrafos vazios
        if not texto_original.strip():
            return

        # 1. Substituir TODOS os placeholders primeiro
        texto_com_placeholders = texto_original
        for placeholder, valor in placeholders.items():
            if placeholder in texto_com_placeholders:
                texto_com_placeholders = texto_com_placeholders.replace(placeholder, valor)

        # 2. Aplicar negrito seletivo
        runs = self.aplicar_negrito_seletivo(texto_com_placeholders, nome_completo)

        # 3. Aplicar ao parágrafo
        paragraph.clear()
        for texto_run, negrito in runs:
            run = paragraph.add_run(texto_run)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)
            run.bold = negrito

    def gerar_contrato(self, template_path, dados_pessoa):
        """Gera contrato com negrito seletivo"""
        try:
            placeholders = self.preparar_placeholders(dados_pessoa)
            nome_completo = dados_pessoa.get('nome_completo', '')

            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template não encontrado: {template_path}")

            print(f"📄 Processando: {os.path.basename(template_path)}")

            # Carregar documento
            doc = Document(template_path)

            # Processar todos os parágrafos
            for paragraph in doc.paragraphs:
                self.processar_paragrafo(paragraph, placeholders, nome_completo)

            # Processar tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.processar_paragrafo(paragraph, placeholders, nome_completo)

            # Nome do arquivo
            nome_base = os.path.splitext(os.path.basename(template_path))[0]
            nome_simplificado = nome_base.replace(" ", "_").replace("_MODEL", "")
            nome_pessoa = nome_completo.replace(" ", "_")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            nome_arquivo = f"{nome_simplificado}_{nome_pessoa}_{timestamp}.docx"
            output_path = os.path.join("output", nome_arquivo)

            # Salvar
            os.makedirs("output", exist_ok=True)
            doc.save(output_path)

            print(f"  ✅ Salvo: {nome_arquivo}")
            return output_path

        except Exception as e:
            print(f"✗ Erro ao gerar contrato: {str(e)}")
            import traceback
            traceback.print_exc()
            raise


def gerar_todos_contratos(dados_pessoa):
    """Função principal"""
    gerador = GeradorContratos()

    templates = [
        ("templates/PROCURACAO_MODEL.docx", "Procuração"),
        ("templates/TERMO DE AUTORIZAÇÃO DE IMAGEM_MODEL.docx", "Autorização de Imagem"),
        ("templates/TERMO DE CONFIDENCIALIDADE_MODEL.docx", "Confidencialidade"),
        ("templates/TERMO DE PROTEÇÃO DE DADOS_MODEL.docx", "Proteção de Dados"),
    ]

    resultados = []

    for template_path, nome_contrato in templates:
        try:
            caminho = gerador.gerar_contrato(template_path, dados_pessoa)
            resultados.append((nome_contrato, caminho))
        except Exception as e:
            print(f"  ✗ Falha no {nome_contrato}: {str(e)}")
            resultados.append((nome_contrato, None))

    return resultados